"""
Word 文档生成服务
将文献综述导出为 .docx 格式
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
from typing import List, Dict


class DocxGenerator:
    def __init__(self):
        pass

    def generate_review_docx(
        self,
        topic: str,
        review: str,
        papers: List[Dict],
        statistics: Dict | None = None
    ) -> bytes:
        """
        生成文献综述的 Word 文档

        Args:
            topic: 论文主题
            review: 综述内容（Markdown 格式）
            papers: 文献列表
            statistics: 统计信息

        Returns:
            Word 文档的二进制内容
        """
        doc = Document()

        # 设置文档样式
        self._setup_document_styles(doc)

        # 添加标题
        title = doc.add_heading(topic, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加摘要/统计信息
        if statistics:
            self._add_statistics(doc, statistics)

        # 解析并添加 Markdown 内容
        self._add_markdown_content(doc, review)

        # 保存到 BytesIO
        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer.read()

    def _setup_document_styles(self, doc):
        """设置文档样式"""
        # 设置正文样式
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)

        # 设置中文字体
        style._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')

    def _add_statistics(self, doc, statistics):
        """添加统计信息"""
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run('文献统计：').bold = True
        stats_text = (
            f"文献总数：{statistics.get('total', 0)}篇 | "
            f"近5年文献占比：{(statistics.get('recent_ratio', 0) * 100):.1f}% | "
            f"英文文献占比：{(statistics.get('english_ratio', 0) * 100):.1f}% | "
            f"总被引量：{statistics.get('total_citations', 0)}次"
        )
        p.add_run(stats_text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_markdown_content(self, doc, review: str):
        """解析并添加 Markdown 内容到文档"""
        lines = review.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i].rstrip()

            # 跳过空行
            if not line or line.isspace():
                i += 1
                continue

            # 标题
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                title_text = line.lstrip('#').strip()
                if title_text:
                    heading = doc.add_heading(title_text, min(level, 6))
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                i += 1
                continue

            # 表格检测
            if '|' in line and line.count('|') >= 2:
                # 可能是表格，先尝试收集表格行
                table_lines = []
                j = i
                while j < len(lines):
                    current_line = lines[j].rstrip()
                    if '|' in current_line and current_line.count('|') >= 2:
                        table_lines.append(current_line)
                        j += 1
                    else:
                        break

                # 检查是否是有效表格
                if len(table_lines) >= 1:
                    # 处理表格，保留所有行（包括分隔线），在 _add_table 中处理
                    self._add_table(doc, table_lines)
                    i = j
                    continue

            # 引用块
            if line.startswith('>'):
                quote_text = line.lstrip('>').strip()
                p = doc.add_paragraph(quote_text, style='Quote')
                i += 1
                continue

            # 无序列表
            if line.strip().startswith(('- ', '* ', '+ ')):
                self._add_list_item(doc, line, ordered=False)
                i += 1
                continue

            # 有序列表
            if re.match(r'^\d+\.\s+', line.strip()):
                self._add_list_item(doc, line, ordered=True)
                i += 1
                continue

            # 分隔线
            if line.strip() in ('---', '***', '___'):
                doc.add_paragraph('_' * 50)
                i += 1
                continue

            # 普通段落（处理行内格式）
            self._add_formatted_paragraph(doc, line)
            i += 1

    def _add_table(self, doc, table_lines: List[str]):
        """添加 Markdown 表格到文档"""
        if not table_lines:
            return

        # 解析每一行，获取单元格数据，同时过滤分隔线
        table_data = []
        header_found = False

        for line in table_lines:
            # 分割单元格，移除首尾的 |
            cells = [cell.strip() for cell in line.split('|')]
            # 移除空的首尾单元格
            if cells and not cells[0]:
                cells = cells[1:]
            if cells and not cells[-1]:
                cells = cells[:-1]

            # 检查是否是分隔线（只包含 -、:、空格的单元格）
            is_separator = False
            if cells:
                all_sep_cells = True
                for cell in cells:
                    cleaned = cell.replace('-', '').replace(':', '').replace(' ', '')
                    if len(cleaned) > 0:
                        all_sep_cells = False
                        break
                if all_sep_cells:
                    is_separator = True

            if not is_separator and cells:
                table_data.append(cells)

        if not table_data:
            return

        # 创建 Word 表格
        rows = len(table_data)
        cols = max(len(row) for row in table_data) if table_data else 0

        if cols == 0:
            return

        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'

        # 填充表格数据
        for row_idx, row_data in enumerate(table_data):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < cols:
                    cell = table.rows[row_idx].cells[col_idx]
                    # 清理单元格文本（移除 Markdown 格式）
                    clean_text = self._strip_markdown_formatting(cell_text)
                    # 添加到单元格
                    paragraph = cell.paragraphs[0]
                    paragraph.text = clean_text
                    # 第一行（表头）加粗
                    if row_idx == 0:
                        for run in paragraph.runs:
                            run.bold = True

    def _add_list_item(self, doc, line, ordered=False):
        """添加列表项"""
        # 移除列表标记
        if ordered:
            text = re.sub(r'^\d+\.\s+', '', line.strip())
        else:
            text = re.sub(r'^[-*+]\s+', '', line.strip())

        # 处理行内格式
        p = doc.add_paragraph(style='List Paragraph')
        self._add_formatted_runs(p, text)
        p.paragraph_format.left_indent = Inches(0.25)

    def _add_formatted_paragraph(self, doc, line):
        """添加带格式的段落"""
        p = doc.add_paragraph()
        self._add_formatted_runs(p, line)

    def _strip_markdown_formatting(self, text: str) -> str:
        """
        移除所有 Markdown 格式符号，只保留纯文本

        Args:
            text: 原始文本（可能包含 Markdown 格式）

        Returns:
            纯文本（移除所有 Markdown 符号）
        """
        # 移除粗体标记 **text** 和 __text__
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'__(.+?)__', r'\1', text)

        # 移除斜体标记 *text* 和 _text_
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = re.sub(r'_(.+?)_', r'\1', text)

        # 移除行内代码标记 `text`
        text = re.sub(r'`(.+?)`', r'\1', text)

        # 移除删除线标记 ~~text~~
        text = re.sub(r'~~(.+?)~~', r'\1', text)

        # 移除未配对的 Markdown 符号（单独出现的 * _ ~ |）
        text = text.replace('*', '')
        text = text.replace('_', '')
        text = text.replace('~', '')
        text = text.replace('|', '')

        # 移除引用标记 [n]
        text = re.sub(r'\[(\d+)\]', r'[\1]', text)  # 保留引用格式

        # 移除链接标记 [text](url)
        text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)

        return text

    def _add_formatted_runs(self, paragraph, text):
        """解析并添加带格式的文本（纯文本模式，不保留 Markdown 格式）"""
        # 直接使用纯文本，移除所有 Markdown 格式符号
        clean_text = self._strip_markdown_formatting(text)

        if clean_text:
            paragraph.add_run(clean_text)
